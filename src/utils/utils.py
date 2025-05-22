import os
import logging
import re
from datetime import datetime
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from contextlib import contextmanager
import functools

timestamped_folder_path = None
# Function to sanitize the scenario name for a valid file name
def sanitize_filename(scenario_name):
    # Replace characters not allowed in file names with underscores
    sanitized_name = re.sub(r'[<>:"/\\|?*]', '_', scenario_name)
    return sanitized_name
# Global variable to hold the single folder path for the test run
# Function to create a single timestamped folder for the test run
def get_or_create_timestamped_folder(base_path, folder_prefix="test_report"):
    global timestamped_folder_path

    if not timestamped_folder_path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        folder_name = f"{folder_prefix}_{timestamp}"
        timestamped_folder_path = os.path.join(base_path, folder_name)

        if not os.path.exists(timestamped_folder_path):
            os.makedirs(timestamped_folder_path)

    return timestamped_folder_path

# Function to generate a unique filename for the screenshot
def generate_unique_filename(scenario_name):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{scenario_name}_{timestamp}.png"

def take_screenshot_with_highlight(self, scenario_name, element_locator=None):
    screenshot_path = capture_and_highlight_screenshot(
        self.driver, 'reports/screenshots', scenario_name, element_locator
    )
    if screenshot_path:
        self.context.screenshots.append(screenshot_path)

#Decorator test
def screenshot_after_action(scenario_name, element_locator=None):
    def decorator(func):
        @functools.wraps(func)
        def wrapper(self, *args, **kwargs):
            result = func(self, *args, **kwargs)
            screenshot_path = capture_and_highlight_screenshot(
                self.driver, 'reports/screenshots', scenario_name, element_locator
            )

            if screenshot_path:
                self.context.screenshots.append(screenshot_path)
            return result
        return wrapper
    return decorator


# Function to capture and highlight a screenshot
def capture_and_highlight_screenshot(driver, base_screenshot_path, message, locator=None):
    try:
        # Reuse the single timestamped folder for the test run
        screenshot_folder = get_or_create_timestamped_folder(base_screenshot_path)

        # Capture the screenshot
        screenshot = driver.get_screenshot_as_png()
        # Generate a unique filename
        screenshot_file_path = os.path.join(screenshot_folder, generate_unique_filename(message))

        # Save the raw screenshot
        with open(screenshot_file_path, 'wb') as file:
            file.write(screenshot)

        # Highlight the element if provided
        if locator:
            screenshot_image = Image.open(screenshot_file_path)
            draw = ImageDraw.Draw(screenshot_image)
            try:
                # Wait for the element to be visible
                element = WebDriverWait(driver, 10).until(
                    EC.visibility_of_element_located(locator)
                )

                # Use JavaScript to get the element's location relative to the document
                element_rect = driver.execute_script("""
                            var rect = arguments[0].getBoundingClientRect();
                            return {
                              x: rect.left,
                              y: rect.top,
                              width: rect.width,
                              height: rect.height,
                              scale: window.devicePixelRatio
                            };
                        """, element)

                # Define the element's bounding box
                scale = element_rect['scale']
                left = element_rect['x'] * scale
                top = element_rect['y'] * scale
                right = (element_rect['x'] + element_rect['width']) * scale
                bottom = (element_rect['y'] + element_rect['height']) * scale

                # Draw a rectangle around the element
                draw.rectangle([left, top, right, bottom], outline="blue", width=5)

                # Save the updated screenshot
                screenshot_image.save(screenshot_file_path)

            except Exception as e:
                logging.warning(f"Could not highlight element: {e}")

        logging.info(f"Screenshot saved at: {screenshot_file_path}")
    except Exception as e:
        logging.error(f"Error capturing or saving screenshot: {e}")
        screenshot_file_path = None

    return screenshot_file_path

# # Function to generate a unique screenshot filename with timestamp
# def generate_unique_filename(scenario_name):
#     timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#     sanitized_name = sanitize_filename(scenario_name)
#     return f"{sanitized_name}_{timestamp}.png"  # Ensure the extension is .png
#
# #Old code for my screenshot
# # Function to capture and highlight a screenshot
# def capture_and_highlight_screenshot(driver, screenshot_path, scenario_name):
#     try:
#         # Capture the screenshot
#         screenshot = driver.get_screenshot_as_png()
#         if not os.path.exists(screenshot_path):
#             os.makedirs(screenshot_path)
#
#         # Generate a unique filename with timestamp
#         screenshot_file_path = os.path.join(screenshot_path, generate_unique_filename(scenario_name))
#
#         # Save the screenshot
#         with open(screenshot_file_path, 'wb') as file:
#             file.write(screenshot)
#
#         # Highlight and save the screenshot
#         screenshot_image = Image.open(screenshot_file_path)
#         draw = ImageDraw.Draw(screenshot_image)
#         #draw.rectangle([100, 200, 500, 300], outline="red", width=5)  # Example region
#         screenshot_image.save(screenshot_file_path)
#
#         logging.info(f"Screenshot saved at: {screenshot_file_path}")
#
#     except Exception as e:
#         logging.error(f"Error capturing or saving screenshot: {e}")
#         screenshot_file_path = None  # Ensure no invalid path is passed
#
#     return screenshot_file_path


# Function to generate the Word document with screenshots
def create_word_document(screenshots, report_path):
    if not screenshots:
        logging.error("No screenshots to add to the Word document.")
        return

    # Get the current date and time for the report header and filename
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")  # Use for unique filename

    # Create the Word document
    doc = Document()

    # Add the timestamp in the title
    doc.add_heading(f'Test Report - Generated on {current_time}', 0)

    # Add the screenshots to the document (only add each screenshot once)
    added_screenshots = set()  # To track added screenshots and avoid duplicates

    for screenshot in screenshots:
        if screenshot and os.path.exists(screenshot):  # Ensure the screenshot exists
            if screenshot not in added_screenshots:  # Check if it's already added
                doc.add_paragraph(f"Screenshot: {screenshot}")
                doc.add_picture(screenshot, width=Inches(5))
                added_screenshots.add(screenshot)  # Mark as added
                logging.info(f"Screenshot added to Word document: {screenshot}")

                try:
                    doc.add_picture(screenshot,width=(5))
                    added_screenshots.add(screenshot)#mark as added
                    logging.info(f"Screenshot added to word document: {screenshot}")
                except Exception as e:
                    logging.info(f"Failed to add screenshot {screenshot} to the word document. Error: {e}")
            else:
                logging.warning(f"Screenshot {screenshot} already added, skipping duplicate.")
        else:
            logging.warning(f"Skipping invalid screenshot path: {screenshot}")

    # Save the Word document in the reports folder with a timestamped filename
    word_report_filename = f"Test_Report_{timestamp}.docx"
    word_report_path = os.path.join(report_path, word_report_filename)

    try:
        doc.save(word_report_path)
        logging.info(f"Word document saved at: {word_report_path}")
    except Exception as e:
        logging.error(f"Error saving Word document: {e}")
